import boto3
import os
import json
import botocore.config
from dotenv import load_dotenv
from fpdf import FPDF
from docx import Document
# loading environment variables
load_dotenv()
# configure Bedrock client
boto3.setup_default_session(profile_name="akaasif+genai-admin")
config = botocore.config.Config(connect_timeout=120, read_timeout=120)
bedrock = boto3.client('bedrock-runtime', 'ap-southeast-2', endpoint_url='https://bedrock-runtime.ap-southeast-2.amazonaws.com',
                       config=config)

# Function to create PDF
class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'Consultancy Services Agreement', 0, 1, 'C')

    def add_table(self, data):
        self.set_font('Arial', '', 12)
        for key, value in data.items():
            self.cell(50, 10, key, 1)
            self.cell(0, 10, value, 1)
            self.ln()

    def add_content(self, content):
        self.set_font('Arial', '', 12)
        self.multi_cell(0, 10, content)
        self.ln()

def create_pdf(details, content):
    print("**** In Create PDF*****")
    pdf = PDF()
    pdf.add_page()
    if(details != ""):
        pdf.add_table(details)
    pdf.add_content(content)
    return pdf

# Function to generate PDF from JSON content
def generate_pdf(data_text, content_type):
    print("**** In Generate PDF*****")
    responde_pdf = PDF()
    if(content_type == 1):
        print('Generate AI content pdf')
        details = {}
        # data_text1 = "# Consultancy Services Agreement\n\n| Principal | Revantage <br> Revantage, Suite 612, Level 2, 150 Castlereagh Street, Melb Vic 2000 |\n|------------|----------------------------------------------------------------------------------|\n| Consultant | John Boggs <br> ABN: Not eligible <br> Address: 1 Bogg Drive, Bogg City, Australia |\n| Commencement Date | 25 March 2026 |\n| Term | As per General Terms clause 1 |\n| Services | As agreed between the parties |\n| Nominated Person | John Boggs |\n| Time Commitment | As required to perform the Services |\n| Principal Representative | As nominated by the Principal from time to time |\n| Required Licences | As required to legally perform the Services |\n| Service Fee | $10,000 |\n| Payment Terms | 14 days from receipt of invoice |\n| Required Insurance | Public liability insurance: $10,000,000 to $90,000,000 <br> Professional indemnity insurance: As required for the Services <br> Workers compensation insurance: As required by law |\n| Required Notice | 1 month |\n\nDated: [Agreement date]\n\nEXECUTED by Revantage by its attorneys under power of attorney dated [date]:\n\n..............................................................\nAttorney  \n\n..............................................................\nName of Attorney (print)\n\n..............................................................\nAttorney\n\n..............................................................  \nName of Attorney (print)  \n\n\nEXECUTED by John Boggs in accordance with section 127(1) of the Corporations Act 2001 (Cth):\n\n..............................................................\nSole Company Secretary and Director\n\n..............................................................\nName of Sole Company Secretary and Director (print)\n\n\n## GENERAL TERMS  \n\n1. **TERM**\n    - **Initial Term**: The Consultant will provide the Services for the Initial Term commencing on the Commencement Date.\n    - **Further Term**: The Principal may, not less than 14 days before the expiry of the Initial Term, elect by written notice to the Consultant to extend the Agreement for a Further Term. The Consultant must inform the Principal in writing of their consent or otherwise to the extension. If agreed, the Consultant will continue to provide the Services under the terms of this Agreement for the Further Term. If not agreed, the Agreement will continue on a month-to-month basis until terminated by either party in accordance with the termination clause.\n\n2. **SERVICES**\n    - **Consultant to Perform Services**: The Principal engages the Consultant to perform the Services. \n    - **Services for Benefit of the Group**: The Consultant will provide the Services for the benefit of the Principal and other companies in the Group or clients of the Principal as directed by the Principal.\n    - **Service Standards**: The Consultant must ensure the Services are provided in accordance with reasonable directions, with due care, skill, and diligence, in compliance with all applicable laws, professional standards, and Revantage's policies and procedures.\n    - **Location of Services**: The Consultant may determine the location from which the Services are performed.\n\n3. **NATURE OF RELATIONSHIP**\n    - **Principal and Independent Consultant**: The Consultant is engaged as an independent non-exclusive contractor and is not an employee of the Principal or any company in the Group.\n    - **No Authority**: The Consultant is not authorized to act on behalf of any company in the Group without the Principal's prior consent.\n\n4. **FEES**\n    - **Service Fee**: Provided the Services are delivered in accordance with this Agreement, the Principal will pay the Consultant the Service Fee within the required timeframe.\n    - **Additional Services**: If additional services are requested and agreed in writing, the Principal will pay the Consultant for the additional services at a rate agreed in writing, exclusive of GST.\n    - **Invoice**: The Consultant must issue an invoice for payment for Services performed on a monthly basis.\n\n5. **EXPENSES**\n    - The Consultant will provide the Services at its own cost and, unless expressly authorized by the Principal, will not be entitled to reimbursement for any out-of-pocket expenses.\n\n6. **EQUIPMENT**\n    - The Consultant must provide all necessary equipment at its expense, including a laptop and mobile phone.  \n\n7. **CONFIDENTIALITY**\n    - **Obligations**: The Consultant must protect the confidential nature of the Confidential Information and not disclose or use it without the Principal's prior written consent, unless required by law.\n    - **Return of Property**: On termination of the Agreement or upon request, the Consultant must return all documents and property of the Principal or the Group.\n\n8. **INTELLECTUAL PROPERTY**\n    - **Ownership**: All Intellectual Property Rights in Contract Materials are owned exclusively by the Principal. The Consultant assigns all such rights to the Principal.\n    - **Third-Party Rights**: The Consultant must not infringe any third-party Intellectual Property Rights in performing the Services and indemnifies the Principal against any claims.\n\n9. **RESPONSIBILITIES OF CONSULTANT**\n    - **Insurance**: The Consultant must maintain the required insurances and provide certificates of currency upon request.\n    - **No Conflict**: The Consultant must not enter into agreements that may conflict with their obligations under this Agreement. \n    - **Indemnity**: The Consultant indemnifies the Principal against liabilities arising from the provision of the Services.\n\n10. **TERMINATION**\n    - This Agreement may be terminated by either party with the Required Notice or immediately for a breach that has not been remedied within 21 days of notice or in cases of bankruptcy or insolvency.\n\n11. **ASSIGNMENT**\n    - The Consultant may not assign their rights under this Agreement without the Principal's written consent. The Principal may assign or subcontract its rights without the Consultant's consent.\n\n12. **PRIVACY AND DATA BREACH**\n    - The Consultant must maintain adequate data protection procedures and notify the Principal immediately of any data breaches, cooperating fully in any investigations.\n\n13. **WORKPLACE HEALTH & SAFETY**\n    - The Consultant must ensure compliance with all applicable workplace health and safety legislation and standards.\n\n14. **OFFICE OF FOREIGN ASSET CONTROL**\n    - The Consultant represents compliance with all anti-terrorism and anti-money laundering regulations and will provide necessary information for the Principal to ensure compliance.\n\n15. **NON-DISCLOSURE OF PRINCIPAL NAME**\n    - The Consultant shall not use the Principal trademarks or name in external publicity without prior written consent.\n\n16. **ANTI-CORRUPTION COMPLIANCE**\n    - The Consultant represents that no payments have been or will be made in violation of anti-corruption laws and agrees to inform the Principal of any violations.\n\n17. **ANTI-MONEY LAUNDERING COMPLIANCE**\n    - The Consultant will comply with all applicable money laundering laws and inform the Principal of any violations.\n\n18. **SANCTIONS**\n    - The Consultant is not subject to any sanctions administered or enforced by relevant authorities and will not knowingly allow a Prohibited Entity to obtain an interest in the Agreement.\n\n19. **GST**\n    - All consideration under this Agreement is exclusive of GST. If GST is payable on any supply, the recipient must pay the amount equal to the GST payable."
        print(data_text)
        lines = data_text.split('\n')
        print(len(lines))
        for line in lines:
            if line.startswith('| '):
                key_value = line.split(' | ')
                if len(key_value) == 2:
                    key = key_value[0].strip().replace('|', '')
                    value = key_value[1].strip().replace('|', '').replace('<br>', '\n')
                    details[key] = value
                    print(f'Key:{key} , value:{value}')
        
        rest_of_content = '\n'.join(lines[16:])
        rest_of_content =  rest_of_content.replace('##','').replace('**','')
        responde_pdf = create_pdf(details,rest_of_content)
    else:
        print('Generate summary pdf')
        data_text = data_text.replace('#','').replace('**','')
        responde_pdf = create_pdf("",data_text)
    return responde_pdf

    

def invoke_llm(bedrock, user_input, doc_template) -> str:
    """
    Creates the initial version of the document based on the details provided by the user.
    :param bedrock: The Amazon Bedrock client that will be used to orchestrate the LLM.
    :param user_input: The details the user is providing to generate the first draft of the document.
    :param doc_template: The document template that the output of the LLM should conform to, to help format and structure
     it more accordingly
    :return: The initial document formatted according to the document template that you pass in with the details provided
    by the user on the front end.
    """

    # Setup Prompt - This prompt passes in the document template and the user input to generate the first draft of the
    # document the user is looking to create
    prompt_data = f"""

Human:

Generate a document based on the user input and the instructions and format provided in the Document Template below  
The tehcnical document should be human readable, well formatted, and broken into the relveant sections.
Response should be in valid Markdown syntax 
###

<Document_Template>
{doc_template}
</Document_Template>
###
<User_Input>
{user_input}
</User_Input>
###

Assistant: Here is a draft based on the provided user input and template

"""
    # Add the prompt to the body to be passed to the Bedrock API along with parameters
    prompt = {
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": 10000,
        "temperature": 0.5,
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": prompt_data
                    }
                ]
            }
        ]
    }
    # formatting the prompt as a json string
    json_prompt = json.dumps(prompt)
    # invoking Claude3, passing in our prompt
    response = bedrock.invoke_model(body=json_prompt, modelId="anthropic.claude-3-sonnet-20240229-v1:0",
                                    accept="application/json", contentType="application/json")
    # getting the response from Claude3 and parsing it to return to the end user
    response_body = json.loads(response.get('body').read())
    print('response_body')
    print(response_body)
    # the final string returned to the end user
    answer = response_body#['content'][0]['text']
    # returning the final string to the end user
    return answer


    """
    This function is specifically focused on refining the document created by invoke_llm, and refining it based on the feedback
    the user is passing in through the frontend.
    :param bedrock: The Amazon Bedrock client that will be used to orchestrate the LLM.
    :param user_feedback: The feedback the user provides through the frontend that contains the addition/changes they would like
    to be made against the original document that was created.
    :param previous_version: This is the original document that was created by the invoke_llm function call.
    :param doc_template: The document template that the output of the LLM should conform to, to help format and structure
     it more accordingly.
    :return: The final version of the document that contains the refinements of the original document specified by the user.
    """
    # Setup Prompt - This prompt passes in the document template and the user feedback, and the previous version to generate the refined draft of the
    # document the user is looking to create.
    prompt_data = f"""

Human:

Refine and Adjust the provided document based on the user feedback and following structure and format guidelines in the Document Template
Response should be in valid Markdown syntax 

###
<document_to_be_refined>
{previous_version}
</document_to_be_refined>

<User_feedback>
{user_feedback}
</User_feedback>

<Document_Template>
{doc_template}
</Document_Template>
###

Assistant: Here is a modified draft press release based on the provided user feedback

"""
    # Add the prompt to the body to be passed to the Bedrock API along with parameters
    prompt = {
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": 1000,
        "temperature": 0.5,
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": prompt_data
                    }
                ]
            }
        ]
    }
    # formatting the prompt as a json string
    json_prompt = json.dumps(prompt)
    # invoking Claude3, passing in our prompt
    response = bedrock.invoke_model(body=json_prompt, modelId="anthropic.claude-3-sonnet-20240229-v1:0",
                                    accept="application/json", contentType="application/json")
    # getting the response from Claude3 and parsing it to return to the end user
    response_body = json.loads(response.get('body').read())
    print('response_body')
    print(response_body)
    # the final string returned to the end user
    answer = response_body['content'][0]['text']
    # returning the final string to the end user
    return answer

def generate_doc(user_input) -> str:
    """
    This function is responsible for orchestrating the call to the invoke_llm function that creates the first version
    of the document.
    :param user_input: The details the user is expecting to be contained in the document.
    :return: The first draft of the document the user is trying to create based on the details provided.
    """
    
    doc_template = """

    Below content should be table formate with 2 columns
    Principal - [Principal Name]
                  [Principal Address]
    Consultant - [Consultant Name]
                  [Consultant Address]
    Commencement Date - [Date of Commencement]
    Term - [Details from template]
    Services - [Services details as in samples or prompt]
    Nominated Person - [Name]
    Time Commitment - [Details from samples or prompt]
    Principal Representative - [Name from prompt and add details from sample]
    Required Licences - [Get details from sample or prompt]
    Service Fee -  [Get details from sample or prompt]
    Payment Terms - [Get details from sample or prompt]
    Required Insurance- [Get insurance details from sample or prompt]
    Required Notice - [Provide notice period from prompt]
    
    #
    Dated:[date]

    EXECUTED by [Principal Name] by its attorneys under power of attorney dated [date]:

    ..............................................................
    Attorney

    ..............................................................
    Name of Attorney (print)
            ..............................................................
    Attorney

    ..............................................................
    Name of Attorney (print)


    EXECUTED by [Consultant Name] in accordance with section 127(1) of the Corporations Act 2001 (Cth):

    ..............................................................
    Sole Company Secretary and Director

    ..............................................................
    Name of Sole Company Secretary and Director
    (print)
	

    GENERAL TERMS

    1. **TERM**
    - **Initial Term**: The Consultant will provide the Services for the Initial Term commencing on the Commencement Date.
    - **Further Term**: The Principal may, not less than 14 days before the expiry of the Initial Term, elect by written notice to the Consultant to extend the Agreement for a Further Term. The Consultant must inform the Principal in writing of their consent or otherwise to the extension. If agreed, the Consultant will continue to provide the Services under the terms of this Agreement for the Further Term. If not agreed, the Agreement will continue on a month-to-month basis until terminated by either party in accordance with the termination clause.

    2. **SERVICES**
    - **Consultant to Perform Services**: The Principal engages the Consultant to perform the Services.
    - **Services for Benefit of the Group**: The Consultant will provide the Services for the benefit of the Principal and other companies in the Group or clients of the Principal as directed by the Principal.
    - **Service Standards**: The Consultant must ensure the Services are provided in accordance with reasonable directions, with due care, skill, and diligence, in compliance with all applicable laws, professional standards, and Revantage’s policies and procedures.
    - **Location of Services**: The Consultant may determine the location from which the Services are performed.

    3. **NATURE OF RELATIONSHIP**
    - **Principal and Independent Consultant**: The Consultant is engaged as an independent non-exclusive contractor and is not an employee of the Principal or any company in the Group.
    - **No Authority**: The Consultant is not authorized to act on behalf of any company in the Group without the Principal’s prior consent.

    4. **FEES**
    - **Service Fee**: Provided the Services are delivered in accordance with this Agreement, the Principal will pay the Consultant the Service Fee within the required timeframe.
    - **Additional Services**: If additional services are requested and agreed in writing, the Principal will pay the Consultant for the additional services at a rate agreed in writing, exclusive of GST.
    - **Invoice**: The Consultant must issue an invoice for payment for Services performed on a monthly basis.

    5. **EXPENSES**
    - The Consultant will provide the Services at its own cost and, unless expressly authorized by the Principal, will not be entitled to reimbursement for any out-of-pocket expenses.

    6. **EQUIPMENT**
    - The Consultant must provide all necessary equipment at its expense, including a laptop and mobile phone.

    7. **CONFIDENTIALITY**
    - **Obligations**: The Consultant must protect the confidential nature of the Confidential Information and not disclose or use it without the Principal’s prior written consent, unless required by law.
    - **Return of Property**: On termination of the Agreement or upon request, the Consultant must return all documents and property of the Principal or the Group.

    8. **INTELLECTUAL PROPERTY**
    - **Ownership**: All Intellectual Property Rights in Contract Materials are owned exclusively by the Principal. The Consultant assigns all such rights to the Principal.
    - **Third-Party Rights**: The Consultant must not infringe any third-party Intellectual Property Rights in performing the Services and indemnifies the Principal against any claims.

    9. **RESPONSIBILITIES OF CONSULTANT**
    - **Insurance**: The Consultant must maintain the required insurances and provide certificates of currency upon request.
    - **No Conflict**: The Consultant must not enter into agreements that may conflict with their obligations under this Agreement.
    - **Indemnity**: The Consultant indemnifies the Principal against liabilities arising from the provision of the Services.

    10. **TERMINATION**
        - This Agreement may be terminated by either party with the Required Notice or immediately for a breach that has not been remedied within 21 days of notice or in cases of bankruptcy or insolvency.

    11. **ASSIGNMENT**
        - The Consultant may not assign their rights under this Agreement without the Principal’s written consent. The Principal may assign or subcontract its rights without the Consultant’s consent.

    12. **PRIVACY AND DATA BREACH**
        - The Consultant must maintain adequate data protection procedures and notify the Principal immediately of any data breaches, cooperating fully in any investigations.

    13. **WORKPLACE HEALTH & SAFETY**
        - The Consultant must ensure compliance with all applicable workplace health and safety legislation and standards.

    14. **OFFICE OF FOREIGN ASSET CONTROL**
        - The Consultant represents compliance with all anti-terrorism and anti-money laundering regulations and will provide necessary information for the Principal to ensure compliance.

    15. **NON-DISCLOSURE OF PRINCIPAL NAME**
        - The Consultant shall not use the Principal trademarks or name in external publicity without prior written consent.

    16. **ANTI-CORRUPTION COMPLIANCE**
        - The Consultant represents that no payments have been or will be made in violation of anti-corruption laws and agrees to inform the Principal of any violations.

    17. **ANTI-MONEY LAUNDERING COMPLIANCE**
        - The Consultant will comply with all applicable money laundering laws and inform the Principal of any violations.

    18. **SANCTIONS**
        - The Consultant is not subject to any sanctions administered or enforced by relevant authorities and will not knowingly allow a Prohibited Entity to obtain an interest in the Agreement.

    19. **GST**
        - All consideration under this Agreement is exclusive of GST. If GST is payable on any supply, the recipient must pay the amount equal to the GST payable.

    """



    # call the invoke_llm function to generate the first draft of the document the user is trying to create
    llmOutput = invoke_llm(bedrock, user_input, doc_template)
    # return the first draft of the document created by the invoke_llm function
    return llmOutput


    """
    This function is specifically responsible for orchestrating the function calls to the invoke_llm_refine to create
    the refined version of the document created.
    :param llm_output: This variable contains the first draft created by the invoke_llm function which is a document created
    according to the details passed in by the user.
    :param user_refine: This is the feedback passed in by the user on the frontend containing the refinements they expect
    to be implemented in the refined version of the document.
    :return: The final version of the document that includes the refinements that the user specified.
    """
    # TODO: EDIT THIS DOCUMENT TEMPLATE TO CONFORM TO THE DOCUMENT FORMAT YOU ARE TRYING TO CREATE
    doc_template = """
        (Press release Style Headline -should be in bold)
        	(Subheader Title: A one sentence summary)
        (LOCATION) - (DATE) - (First paragraph: summary of the service/product/feature "launch")
        (Second Paragraph: The second paragraph explains the opportunity or problem that needs to be solved)
        (The third paragraph gives the approach or the solution.)
        (The fourth paragraph quotes an Internal leader.)
        (The fifth paragraph describes the customer experience - how users will discover and use what you propose)
        (The sixth paragraph includes a specific, believable, human-sounding customer testimonial.)
        (The seventh paragraph directs the reader where to go to get started)
        #

        Customer FAQ's 
        [This section will consist of questions and answers relevant to customers and user]
        1. [Question]
        	A: [Answer to Question]
        2. [Question]
        	A: [Answer to Question]
        …
        X. [Question]
        	A: [Answer to Question]

        StakeHolder FAQ's 
        [This section will consist of questions and answers relevant to customers and user]
        1. [Question - Should be professional and concise]
        	A: [Answer to Question - should be professional and relevant]
        2. [Question]
        	A: [Answer to Question]
        …
        X. [Question]
        	A: [Answer to Question]

        Appendices
        [If you used specific data points or references in the section describing your approach, include more complete data set as an appendix. Add relevant well sourced data points and studies]
        Appendix A: (Studies, statistics and Supporting evidence reference material directly relevant to your Press Release - should be detailed, relevant and well sourced)

        Appendix B: (Studies, statistics and Supporting evidence reference material directly relevant to your Press Release - should be detailed, relevant and well sourced)
        …
        Appendix X: (Studies, statistics and Supporting evidence reference material directly relevant to your Press Release - should be detailed, relevant and well sourced)
        """
    # call the invoke_llm_refine function to generate the refined draft of the document that the user is trying to create
    llmOutput = invoke_llm_refine(bedrock,user_refine, llm_output, doc_template)
    # return the refined version of the document create by the invoke_llm_refine function
    return llmOutput

def summarize_read_docx(file):
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def summarize_process_file_content(content):
    prompt = "Generate summary for the content, creat ten bullet point summary"
    user_input = "\n".join([prompt, content])
    # print(user_input)
    doc_template = """
        [Document Title]

        [Provide the main heading or title of the document here]
        ## Quick Summary
        [Provide a brief summary or abstract of the document's content here. This should be a short paragraph that gives an overview of the main points or purpose of the document.]

        ## Summarized Bullet Points
        - **Key Point 1**: [Provide a concise description of the first key point here]
        - **Key Point 2**: [Provide a concise description of the second key point here]
        - **Key Point 3**: [Provide a concise description of the third key point here]
        - **Key Point 4**: [Provide a concise description of the fourth key point here]
        - **Key Point 5**: [Provide a concise description of the fifth key point here]
        …
        - **Key Point X**: [Provide a concise description of the fifth key point here]
    
    """
    # call the invoke_llm function to generate the first draft of the document the user is trying to create
    llmOutput = invoke_llm(bedrock, user_input, doc_template)
    # return the first draft of the document created by the invoke_llm function
    print(f'************LLM Output :************ \n {llmOutput}')
    return llmOutput